"""DOCX extraction: structure, typography, layout, formatting and visuals."""

from __future__ import annotations

import io
import re
from collections import Counter
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Length

from ..models import (
    DocumentProfile,
    Formatting,
    Heading,
    Layout,
    Typography,
    Visual,
)

ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justified",
}

PAGE_SIZES_IN = {
    "A4": (8.27, 11.69),
    "Letter": (8.5, 11.0),
    "Legal": (8.5, 14.0),
    "A5": (5.83, 8.27),
}

NUMBERED_HEADING = re.compile(r"^\s*\d+(\.\d+)*[.)]?\s+\S")
CAPTION = re.compile(r"^\s*(figure|fig\.|table)\s*\d+", re.I)


def _inches(value: Optional[Length]) -> Optional[float]:
    return round(value.inches, 2) if value is not None else None


def _page_size(width_in: Optional[float], height_in: Optional[float]) -> Optional[str]:
    if not width_in or not height_in:
        return None
    for name, (width, height) in PAGE_SIZES_IN.items():
        if abs(width - width_in) < 0.2 and abs(height - height_in) < 0.3:
            return name
    return f'{width_in}" x {height_in}"'


def _heading_level(style_name: str) -> Optional[int]:
    match = re.match(r"heading\s*(\d)", style_name.strip(), re.I)
    if match:
        return int(match.group(1))
    if style_name.strip().lower() in {"title", "subtitle"}:
        return 1 if style_name.strip().lower() == "title" else 2
    return None


def extract_docx(data: bytes, file_name: str) -> DocumentProfile:
    document = Document(io.BytesIO(data))

    headings: List[Heading] = []
    body_fonts: Counter[str] = Counter()
    body_sizes: Counter[float] = Counter()
    body_alignments: Counter[str] = Counter()
    line_spacings: Counter[float] = Counter()
    heading_fonts: Dict[str, str] = {}
    heading_sizes: Dict[str, float] = {}
    heading_bold: List[bool] = []
    heading_colors: Counter[str] = Counter()
    list_styles: set[str] = set()
    caption_positions: List[str] = []
    numbered_headings = 0
    sections: Dict[str, str] = {}
    current_section = "Preamble"
    section_text: Dict[str, List[str]] = {current_section: []}
    all_text: List[str] = []
    order = 0
    previous_was_image = False

    default_font = document.styles["Normal"].font
    default_font_name = default_font.name
    default_size = default_font.size.pt if default_font.size else None

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        has_image = bool(paragraph._p.findall(".//{*}blip"))

        if text:
            all_text.append(text)

        level = _heading_level(style_name)
        if level and text:
            order += 1
            headings.append(Heading(text=text, level=level, order=order))
            key = f"Heading {level}"
            run = paragraph.runs[0] if paragraph.runs else None
            if run is not None:
                if run.font.name or default_font_name:
                    heading_fonts.setdefault(key, run.font.name or default_font_name or "")
                if run.font.size:
                    heading_sizes.setdefault(key, round(run.font.size.pt, 1))
                heading_bold.append(bool(run.bold))
                color = getattr(run.font.color, "rgb", None)
                if color:
                    heading_colors[f"#{color}"] += 1
            if NUMBERED_HEADING.match(text):
                numbered_headings += 1
            current_section = text
            section_text.setdefault(current_section, [])
            previous_was_image = has_image
            continue

        if "list" in style_name.lower():
            list_styles.add(style_name)

        if text and CAPTION.match(text):
            caption_positions.append("below" if previous_was_image else "above")

        if text:
            section_text.setdefault(current_section, []).append(text)
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                body_fonts[run.font.name or default_font_name or "Unknown"] += 1
                size = run.font.size.pt if run.font.size else default_size
                if size:
                    body_sizes[round(size, 1)] += 1
            alignment = ALIGNMENT_NAMES.get(paragraph.alignment)
            if alignment is None and paragraph.style is not None:
                alignment = ALIGNMENT_NAMES.get(
                    paragraph.style.paragraph_format.alignment
                )
            body_alignments[alignment or "left"] += 1
            spacing = paragraph.paragraph_format.line_spacing
            if isinstance(spacing, float):
                line_spacings[round(spacing, 2)] += 1

        previous_was_image = has_image

    section = document.sections[0] if document.sections else None
    width_in = _inches(section.page_width) if section else None
    height_in = _inches(section.page_height) if section else None
    margins = {}
    if section:
        margins = {
            "top": _inches(section.top_margin) or 0,
            "bottom": _inches(section.bottom_margin) or 0,
            "left": _inches(section.left_margin) or 0,
            "right": _inches(section.right_margin) or 0,
        }

    image_count = len(document.inline_shapes)

    profile = DocumentProfile(
        file_name=file_name,
        source_type="docx",
        headings=headings,
        typography=Typography(
            body_font=body_fonts.most_common(1)[0][0] if body_fonts else default_font_name,
            body_size_pt=body_sizes.most_common(1)[0][0] if body_sizes else default_size,
            heading_fonts=heading_fonts,
            heading_sizes_pt=heading_sizes,
            bold_headings=(sum(heading_bold) > len(heading_bold) / 2)
            if heading_bold
            else None,
        ),
        layout=Layout(
            page_size=_page_size(width_in, height_in),
            margins_in=margins,
            body_alignment=body_alignments.most_common(1)[0][0]
            if body_alignments
            else None,
            line_spacing=line_spacings.most_common(1)[0][0] if line_spacings else None,
        ),
        formatting=Formatting(
            heading_numbering=numbered_headings > max(1, len(headings) // 2),
            list_styles=sorted(list_styles),
            table_count=len(document.tables),
            caption_position=Counter(caption_positions).most_common(1)[0][0]
            if caption_positions
            else None,
        ),
        visual=Visual(
            image_count=image_count,
            image_alignment=None,
            heading_color=heading_colors.most_common(1)[0][0] if heading_colors else None,
            accent_colors=[color for color, _ in heading_colors.most_common(3)],
        ),
        sections={
            name: "\n".join(lines)[:4000] for name, lines in section_text.items() if lines
        },
        terminology=_terminology(all_text),
        word_count=sum(len(line.split()) for line in all_text),
        text_excerpt="\n".join(all_text)[:12000],
    )
    return profile


def _terminology(lines: List[str]) -> List[str]:
    """Capitalised multi-word terms that look like defined terminology."""
    counter: Counter[str] = Counter()
    pattern = re.compile(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b")
    for line in lines:
        for match in pattern.findall(line):
            counter[match] += 1
    return [term for term, count in counter.most_common(20) if count > 1]
